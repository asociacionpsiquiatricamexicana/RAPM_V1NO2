#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""manuscrito_desde_md.py — arma un manuscrito .docx desde un Markdown.

El puente entre "el material llego en cualquier formato" y el receptor.
taller/recibir_articulo.py sabe leer .docx; este script produce ese .docx
desde un Markdown, que es lo que queda cuando alguien (persona o modelo)
reconstruye un manuscrito a partir de un PDF, de fotos de hojas, de un
dictado ya transcrito o de un monton de datos sueltos.

Por que pasar por un archivo y no armar el .docx a mano cada vez: porque el
repositorio exige que lo entregado se pueda regenerar desde lo versionado.
El .md queda junto al articulo como fuente de la reconstruccion, y quien
quiera auditar de donde salio cada parrafo tiene el archivo, no el recuerdo
de una conversacion.

NO inventa contenido. Lo que no este en el .md no aparece en el .docx; los
huecos se resuelven ANTES, preguntandole a quien entrega el material.

Formato de entrada (el que produce naturalmente quien transcribe):

    # Titulo del articulo
    Autor Uno, Autor Dos y Autor Tres
    Afiliacion institucional, Ciudad, Pais
    correo@ejemplo.mx
    ORCID: 0000-0000-0000-0000

    ## Resumen
    Texto del resumen...

    **Palabras clave:** una, dos, tres

    ## Introduccion
    Cuerpo...

    ### Subseccion
    Cuerpo...

    Tabla 1. Pie de la tabla.
    | Columna | Otra |
    | --- | --- |
    | dato | dato |

    Figura 1. Pie de la figura.
    ![](ruta/a/imagen.png)

    ## Referencias
    Apellido, N. (2024). Titulo. Revista, 1(1), 1-10.

Uso:
    python3 taller/manuscrito_desde_md.py FUENTE.md [-o SALIDA.docx]
"""
import argparse
import os
import re
import sys

try:
    import docx
    from docx.shared import Inches
except ImportError:
    sys.exit(
        "falta python-docx: python3 -m pip install python-docx "
        "(vease taller/sondas/requisitos.txt)"
    )

RE_FILA_TABLA = re.compile(r"^\s*\|(.+)\|\s*$")
RE_SEPARADOR_TABLA = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
RE_IMAGEN = re.compile(r"^\s*!\[[^\]]*\]\(([^)]+)\)\s*$")
RE_ENFASIS = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|_(.+?)_")


def celdas_de_fila(linea):
    interior = RE_FILA_TABLA.match(linea).group(1)
    return [c.strip() for c in interior.split("|")]


def agrega_parrafo_con_enfasis(doc, texto, estilo=None):
    """Conserva **negrita** y *cursiva* como formato real del .docx, para que
    recibir_articulo.py los traduzca a \\textbf/\\textit. El resto va tal cual:
    este script no reescribe el texto, solo lo coloca."""
    p = doc.add_paragraph(style=estilo)
    pos = 0
    for m in RE_ENFASIS.finditer(texto):
        if m.start() > pos:
            p.add_run(texto[pos:m.start()])
        negrita = m.group(1) is not None or m.group(3) is not None
        contenido = next(g for g in m.groups() if g is not None)
        run = p.add_run(contenido)
        run.bold = negrita
        run.italic = not negrita
        pos = m.end()
    if pos < len(texto):
        p.add_run(texto[pos:])
    return p


def _sin_acentos_min(texto):
    import unicodedata
    s = unicodedata.normalize("NFKD", texto.strip().lower())
    return "".join(c for c in s if not unicodedata.combining(c))


def _es_seccion_referencias(texto):
    n = _sin_acentos_min(texto)
    return "referencias" in n or "bibliografia" in n or "references" in n


def construir(ruta_md, ruta_salida):
    base_md = os.path.dirname(os.path.abspath(ruta_md))
    with open(ruta_md, encoding="utf-8") as fh:
        lineas = fh.read().splitlines()

    doc = docx.Document()
    i = 0
    faltantes = []
    # Dos zonas donde cada renglon es una unidad y NO debe unirse con el
    # siguiente, aunque Markdown normalmente los juntaria:
    #   - el bloque de cabecera (antes del primer "##"): autor, afiliacion,
    #     correo y ORCID son campos distintos, y recibir_articulo.py los
    #     distingue renglon por renglon. Unirlos metia los cuatro dentro de
    #     \APMauthor y dejaba afiliacion/correo/ORCID en [PENDIENTE].
    #   - la lista de referencias: cada entrada es una referencia, y unirlas
    #     produce una sola \APMref con todas dentro.
    # En el cuerpo si se unen, que es donde la prosa viene con renglones
    # partidos y debe fluir como un parrafo.
    en_cabecera = True
    en_referencias = False

    while i < len(lineas):
        linea = lineas[i]
        crudo = linea.strip()

        if not crudo:
            i += 1
            continue

        # Encabezados: # -> Title, ## -> Heading 1, ### -> Heading 2
        m = re.match(r"^(#{1,6})\s+(.*)$", crudo)
        if m:
            nivel = len(m.group(1))
            texto = m.group(2).strip()
            if nivel == 1:
                doc.add_heading(texto, level=0)
            else:
                doc.add_heading(texto, level=min(nivel - 1, 4))
                en_cabecera = False
                en_referencias = _es_seccion_referencias(texto)
            i += 1
            continue

        # Imagen: ![](ruta)
        m = RE_IMAGEN.match(linea)
        if m:
            ruta_img = m.group(1).strip()
            if not os.path.isabs(ruta_img):
                ruta_img = os.path.join(base_md, ruta_img)
            if os.path.exists(ruta_img):
                doc.add_picture(ruta_img, width=Inches(5.4))
            else:
                # No se sustituye por texto ni se omite en silencio: se
                # declara, para que quien entrego el material la aporte.
                faltantes.append(ruta_img)
                doc.add_paragraph(f"[FALTA LA IMAGEN: {m.group(1).strip()}]")
            i += 1
            continue

        # Tabla en tuberias
        if RE_FILA_TABLA.match(linea):
            filas = []
            while i < len(lineas) and RE_FILA_TABLA.match(lineas[i]):
                if not RE_SEPARADOR_TABLA.match(lineas[i]):
                    filas.append(celdas_de_fila(lineas[i]))
                i += 1
            if filas:
                ncols = max(len(f) for f in filas)
                tabla = doc.add_table(rows=0, cols=ncols)
                tabla.style = "Table Grid"
                for fila in filas:
                    celdas = tabla.add_row().cells
                    for j in range(ncols):
                        celdas[j].text = fila[j] if j < len(fila) else ""
            continue

        # Parrafo. En cabecera y en referencias, un renglon = un parrafo.
        # En el cuerpo se unen los renglones contiguos, como Markdown, para
        # no partir la prosa del autor en renglones sueltos.
        bloque = [crudo]
        i += 1
        if not (en_cabecera or en_referencias):
            while i < len(lineas):
                siguiente = lineas[i].strip()
                if (not siguiente or siguiente.startswith("#")
                        or RE_FILA_TABLA.match(lineas[i]) or RE_IMAGEN.match(lineas[i])):
                    break
                bloque.append(siguiente)
                i += 1
        agrega_parrafo_con_enfasis(doc, " ".join(bloque))

    doc.save(ruta_salida)
    return faltantes


def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Arma un manuscrito .docx desde un Markdown, para "
                    "alimentarlo a taller/recibir_articulo.py."
    )
    ap.add_argument("fuente", help="ruta al .md reconstruido")
    ap.add_argument("-o", "--salida", default=None,
                    help="ruta del .docx (por omision, junto al .md)")
    args = ap.parse_args(argv)

    if not os.path.exists(args.fuente):
        sys.exit(f"no existe la fuente: {args.fuente}")
    salida = args.salida or os.path.splitext(args.fuente)[0] + ".docx"

    faltantes = construir(args.fuente, salida)
    print(f"escrito: {salida}")
    if faltantes:
        print("IMAGENES QUE NO SE ENCONTRARON (pidelas antes de componer):",
              file=sys.stderr)
        for f in faltantes:
            print(f"  - {f}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
