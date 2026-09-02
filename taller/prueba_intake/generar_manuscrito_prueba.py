#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""generar_manuscrito_prueba.py — construye manuscrito_prueba.docx.

ESTO NO ES UN MANUSCRITO REAL. Es un fixture sintetico de regresion para
taller/recibir_articulo.py (analogo a taller/ejemplo_articulo_original.tex
para el taller de composicion, pero para el receptor de manuscritos .docx).
Autores, afiliaciones, ORCID, DOI, cifras y referencias son INVENTADOS y no
deben confundirse con investigacion real, citarse, ni reutilizarse como
material editorial. El dominio "ejemplo.invalid" en URLs y correos esta
reservado para pruebas (RFC 2606) — a proposito, para que nadie lo confunda
con una direccion real.

Uso: python3 taller/prueba_intake/generar_manuscrito_prueba.py
Produce: taller/prueba_intake/manuscrito_prueba.docx
"""
import os

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw

AQUI = os.path.dirname(os.path.abspath(__file__))
TALLER = os.path.dirname(AQUI)
SALIDA = os.path.join(AQUI, "manuscrito_prueba.docx")
IMAGEN_PRUEBA = os.path.join(AQUI, "figura_prueba.png")


def genera_imagen_prueba():
    """PNG simple (no el logo de alta resolucion, que pesa >300KB y hace
    que el PDF exceda los 600KB del despliegue): un rectangulo con texto
    de relleno, suficiente para ejercitar el flotante de figura."""
    ancho, alto = 900, 500
    img = Image.new("RGB", (ancho, alto), (245, 240, 232))
    dibujo = ImageDraw.Draw(img)
    dibujo.rectangle([20, 20, ancho - 20, alto - 20], outline=(120, 40, 45), width=6)
    for i in range(6):
        y = 90 + i * 60
        dibujo.line([(60, y), (ancho - 60, y)], fill=(90, 90, 90), width=3)
    dibujo.ellipse([ancho / 2 - 80, alto / 2 - 80, ancho / 2 + 80, alto / 2 + 80],
                    outline=(120, 40, 45), width=5)
    img.save(IMAGEN_PRUEBA, format="PNG", optimize=True)
    return IMAGEN_PRUEBA


def parrafo_normal(doc, texto):
    return doc.add_paragraph(texto)


def main():
    doc = Document()

    # ── Portadilla ───────────────────────────────────────────
    doc.add_heading(
        "Acompañamiento entre pares y bienestar percibido en estudiantes "
        "de posgrado: un estudio piloto de fixture",
        level=0,
    )
    parrafo_normal(
        doc,
        "Ana Prueba Fixture, Beatriz Ejemplo Ficticia y Carlos Simulado "
        "Ficticio",
    )
    parrafo_normal(
        doc,
        "Programa de Fixtures Editoriales, Instituto Imaginario de Pruebas "
        "Taller, Ciudad de México, México",
    )
    parrafo_normal(doc, "ana.prueba@ejemplo.invalid")
    parrafo_normal(doc, "ORCID: 0000-0001-2345-6789")
    parrafo_normal(doc, "Tel.: +52 55 0001 0002")

    # ── Resumen ──────────────────────────────────────────────
    doc.add_heading("Resumen", level=2)
    parrafo_normal(
        doc,
        "Introducción. Este resumen forma parte de un archivo de prueba "
        "generado por taller/prueba_intake/generar_manuscrito_prueba.py "
        "y no describe un estudio real. Método. Se simuló un diseño "
        "descriptivo con datos ficticios de acompañamiento entre pares. "
        "Resultados. Los valores mostrados en la Tabla 1 son inventados "
        "con fines de prueba tipográfica. Conclusiones. El propósito único "
        "de este documento es ejercitar la extracción automatizada de "
        "taller/recibir_articulo.py.",
    )
    parrafo_normal(
        doc,
        "Palabras clave: Fixture de prueba, Extracción automatizada, "
        "Acompañamiento entre pares, Bienestar percibido, Datos ficticios",
    )

    # ── Cuerpo IMRaD ─────────────────────────────────────────
    doc.add_heading("Introducción", level=1)
    parrafo_normal(
        doc,
        "Este texto de introducción es un relleno deliberado, escrito para "
        "ejercitar el receptor de manuscritos con acentos (á, é, í, ó, ú, "
        "ñ, Ñ, ü) y con formato de énfasis en línea.",
    )
    p = doc.add_paragraph()
    p.add_run("Esta oración combina ")
    r_b = p.add_run("una palabra en negrita")
    r_b.bold = True
    p.add_run(", ")
    r_i = p.add_run("una palabra en cursiva")
    r_i.italic = True
    p.add_run(" y texto sin formato, para verificar que run.bold/run.italic "
               "se traducen a \\textbf{}/\\textit{}.")

    doc.add_heading("Método", level=1)
    doc.add_heading("Participantes", level=2)
    parrafo_normal(
        doc,
        "Participantes ficticios de un archivo de prueba: no representan "
        "personas reales ni datos clínicos.",
    )
    doc.add_heading("Procedimiento", level=2)
    parrafo_normal(
        doc,
        "El procedimiento descrito aquí es inventado con el único fin de "
        "poblar la sección Método con dos subsecciones, tal como exige el "
        "fixture de regresión.",
    )

    doc.add_heading("Resultados", level=1)
    parrafo_normal(
        doc,
        "La Tabla 1 resume valores ficticios organizados en cuatro "
        "columnas y cuatro filas, incluida la fila de encabezado.",
    )
    parrafo_normal(doc, "Tabla 1. Valores ficticios de prueba tipográfica, sin significado clínico.")
    tabla = doc.add_table(rows=4, cols=4)
    tabla.style = "Table Grid"
    encabezados = ["Grupo", "n (ficticio)", "Media (ficticia)", "DE (ficticia)"]
    filas = [
        ["Acompañamiento", "18", "4.2", "0.9"],
        ["Control", "17", "3.1", "1.1"],
        ["Total", "35", "3.7", "1.1"],
    ]
    for j, texto in enumerate(encabezados):
        tabla.cell(0, j).text = texto
    for i, fila in enumerate(filas, start=1):
        for j, texto in enumerate(fila):
            tabla.cell(i, j).text = texto
    parrafo_normal(
        doc,
        "Tras la tabla, este párrafo continúa el cuerpo para confirmar que "
        "el receptor conserva el orden real de intercalado entre texto y "
        "tabla.",
    )

    doc.add_heading("Discusión", level=1)
    parrafo_normal(
        doc,
        "Esta discusión ficticia sirve únicamente para ejercitar la "
        "extracción de un encabezado H1 más entre Resultados y "
        "Conclusiones.",
    )

    doc.add_heading("Conclusiones", level=1)
    parrafo_normal(
        doc,
        "La Figura 1 ilustra un elemento gráfico incrustado en su posición "
        "real dentro del cuerpo del documento.",
    )
    parrafo_normal(doc, "Figura 1. Marca de prueba incluida únicamente para ejercitar el flotante de figura.")
    ruta_imagen = genera_imagen_prueba()
    doc.add_picture(ruta_imagen, width=Inches(4.6))
    parrafo_normal(
        doc,
        "Este párrafo final de las conclusiones confirma que el cuerpo "
        "continúa después de la figura sin perder texto.",
    )

    parrafo_normal(
        doc,
        "Conflicto de intereses: Los autores de este archivo de prueba "
        "declaran no tener conflicto de intereses; el contenido es "
        "ficticio y no corresponde a un estudio real.",
    )

    # ── Referencias (literal, con acentos y URLs largas) ────
    doc.add_heading("Referencias", level=1)
    referencias = [
        "Prueba Fixture, A., & Ejemplo Ficticia, B. (2026). Generación "
        "automatizada de artículos de prueba para talleres editoriales: "
        "un ejercicio de composición tipográfica. Revista Imaginaria de "
        "Pruebas de Taller, 1(1), 1–9. "
        "https://ejemplo.invalid/doi/10.9999/rapm.taller.000001",
        "Simulado Ficticio, C. (2025). Acompañamiento entre pares como "
        "constructo simulado: notas metodológicas de un fixture. "
        "Cuadernos Ficticios de Metodología, 4(2), 55–61. "
        "https://ejemplo.invalid/articulos/simulado-ficticio-2025-notas-metodologicas-fixture-taller",
        "Ejemplo Ficticia, B., Prueba Fixture, A., & Datos Inventados, D. "
        "(2024). Diseño de datos ficticios con acentuación variada: "
        "íconos, diéresis y eñes en referencias bibliográficas de "
        "prueba. Anales Simulados de Composición, 12(3), 300–318. "
        "https://ejemplo.invalid/doi/10.9999/rapm.taller.000002",
        "Falso Autor, E. (2023). Notas sobre la extracción de tablas y "
        "figuras incrustadas en documentos .docx: un memorando interno "
        "de prueba. Boletín Interno Ficticio, 7, 12–20.",
        "Ficticio Ñáñez, Í. (2022). Ñoñerías tipográficas: cómo probar "
        "acentos y eñes en un aparato bibliográfico simulado. Revista "
        "Inventada de Ortotipografía, 9(1), 44–50. "
        "https://ejemplo.invalid/doi/10.9999/rapm.taller.000003",
        "Comité Editorial Imaginario. (2026). Lineamientos internos "
        "(ficticios) para la recepción automatizada de manuscritos. "
        "Documento de trabajo no publicado del taller.",
        "Datos Inventados, D., & Falso Autor, E. (2021). Un panorama "
        "simulado de veinte años de artículos de prueba en revistas "
        "imaginarias de habla hispana, con especial atención a la "
        "extensión de las URL. Serie Ficticia de Reportes Técnicos, 3, "
        "1–5. https://ejemplo.invalid/reportes/panorama-simulado-veinte-anos-articulos-prueba-revistas-imaginarias-habla-hispana",
        "Organización Mundial Ficticia de Pruebas Editoriales. (2020). "
        "Recomendaciones (inventadas) para la composición tipográfica de "
        "manuscritos electrónicos. https://ejemplo.invalid/recomendaciones/2020",
        "Prueba Fixture, A. (2019). Diez años de fixtures: una revisión "
        "simulada no sistemática de archivos de prueba tipográfica. "
        "Revista Imaginaria de Pruebas de Taller, 3(4), 210–225. "
        "https://ejemplo.invalid/doi/10.9999/rapm.taller.000004",
        "Ejemplo Ficticia, B. (2018). El acento y la eñe en la "
        "composición LaTeX de pruebas automatizadas: álgebra, épica y "
        "ortografía simuladas. Cuadernos Ficticios de Metodología, 1(1), "
        "1–3.",
        "Simulado Ficticio, C., & Ficticio Ñáñez, Í. (2017). Comparación "
        "simulada entre dos flujos de composición tipográfica ficticios, "
        "con notas sobre desbordamiento de cajas. Anales Simulados de "
        "Composición, 5(2), 88–104. "
        "https://ejemplo.invalid/doi/10.9999/rapm.taller.000005",
    ]
    for r in referencias:
        parrafo_normal(doc, r)

    doc.save(SALIDA)
    print(f"escrito: {SALIDA}")


if __name__ == "__main__":
    main()
