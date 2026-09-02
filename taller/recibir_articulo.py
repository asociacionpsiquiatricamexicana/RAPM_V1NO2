#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""recibir_articulo.py — recibe un manuscrito .docx y produce el articulo
camera-ready de la Revista de la Asociacion Psiquiatrica Mexicana, A.C.,
compilado con apm-editorial.cls, sin que nadie escriba LaTeX a mano.

Uso:
    python3 taller/recibir_articulo.py MANUSCRITO.docx
        [--tipo "Articulo original"]
        [--numero VOL6_NO3]
        [--raiz numeros/]
        [--art N]

Sigue numeros/LEEME.md al pie de la letra para la ubicacion y el nombre de
los archivos de salida. No toca el texto de los autores: solo escapa
caracteres especiales de LaTeX: nunca reescribe, reordena ni corrige.
Cuando falta un dato editorial (fechas de aceptacion/publicacion, DOI,
ORCID, conflicto de intereses, financiamiento, pies de tabla/figura...) se
inserta un placeholder marcado "[PENDIENTE: ...]"; nunca se inventa.
"""
from __future__ import annotations

import argparse
import datetime
import os
import re
import shutil
import subprocess
import sys
import unicodedata

try:
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    sys.exit(
        "falta python-docx: python3 -m pip install python-docx "
        "(vease taller/sondas/requisitos.txt)"
    )

AQUI = os.path.dirname(os.path.abspath(__file__))
RAIZ_REPO = os.path.dirname(AQUI)

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

ARTICULOS_MINUSCULOS = {"el", "la", "los", "las", "un", "una", "de"}

MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

# tabla del enunciado: acentuadas -> octal (PDFDocEncoding / Latin-1)
OCTAL_ACENTOS = {
    "á": "341", "é": "351", "í": "355", "ó": "363", "ú": "372",
    "ñ": "361", "ü": "374",
    "Á": "301", "É": "311", "Í": "315", "Ó": "323", "Ú": "332",
    "Ñ": "321", "Ü": "334",
}

RE_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
RE_ORCID = re.compile(r"\d{4}-\d{4}-\d{4}-\d{3}[\dX]")
RE_DOI_BARE = re.compile(r"10\.\d{4,9}/\S+")
RE_DOI_URL = re.compile(r"https?://(?:dx\.)?doi\.org/\S+", re.IGNORECASE)
RE_TEL_CONTEXTO = re.compile(r"tel[eé]?fono|\btel\.?\b|\bphone\b", re.IGNORECASE)
RE_TEL_NUM = re.compile(r"[+]?[\d][\d\s().-]{5,}\d")
RE_URL = re.compile(r"(?:https?://|www\.)\S+", re.IGNORECASE)


# ═══════════════════════════════════════════════════════════════
# Texto: normalizacion, escapado LaTeX, PDF metadata en octal
# ═══════════════════════════════════════════════════════════════

def normaliza(texto: str) -> str:
    """minusculas, sin acentos: para comparar rotulos de encabezado."""
    if not texto:
        return ""
    s = texto.strip().lower()
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c))


_MARCA_BACKSLASH = "\x00BACKSLASH\x00"


def escapar_latex(texto: str) -> str:
    """Escapa TODO texto extraido del docx antes de insertarlo en el .tex.
    Backslash primero, luego & % $ # _ { } ~ ^ (orden del enunciado). El
    backslash se sustituye por una marca temporal (no por \\textbackslash{}
    todavia): si se sustituyera de una vez, las llaves de \\textbackslash{}
    quedarian expuestas al paso siguiente y este las volveria a escapar
    (\\textbackslash\\{\\}), rompiendo el comando. La marca se resuelve
    al final, cuando ya no queda ningun paso de escapado por correr."""
    if texto is None:
        return ""
    resultado = texto.replace("\\", _MARCA_BACKSLASH)
    for viejo, nuevo in (
        ("&", r"\&"), ("%", r"\%"), ("$", r"\$"), ("#", r"\#"),
        ("_", r"\_"), ("{", r"\{"), ("}", r"\}"),
        ("~", r"\textasciitilde{}"), ("^", r"\textasciicircum{}"),
    ):
        resultado = resultado.replace(viejo, nuevo)
    return resultado.replace(_MARCA_BACKSLASH, r"\textbackslash{}")


def a_octal_pdfinfo(texto: str) -> str:
    """Traduce cada caracter acentuado a su octal \\NNN (via \\string), deja
    el resto tal cual. Replica el patron de taller/ejemplo_editorial.tex.
    Para /Author y /Keywords de \\pdfinfo, que NO soportan UTF-8 directo."""
    if not texto:
        return ""
    partes = []
    for ch in texto:
        if ch in OCTAL_ACENTOS:
            partes.append("\\string\\" + OCTAL_ACENTOS[ch])
        else:
            partes.append(ch)
    return "".join(partes)


def sanea_nombre_archivo(palabra: str) -> str:
    """Mayusculas, sin acentos/ene, solo [A-Z0-9] — nunca texto crudo del
    docx como ruta (nada de '..', separadores, ni nombres kilometricos)."""
    s = unicodedata.normalize("NFKD", palabra or "")
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"[^A-Za-z0-9]", "", s).upper()
    return s[:40]


def primera_palabra_significativa(titulo: str) -> str:
    palabras = re.findall(r"[^\W\d_]+", titulo or "", flags=re.UNICODE)
    for p in palabras:
        if normaliza(p) not in ARTICULOS_MINUSCULOS:
            saneada = sanea_nombre_archivo(p)
            if saneada:
                return saneada
    return "ARTICULO"


def fecha_hoy_es() -> str:
    hoy = datetime.date.today()
    return f"{hoy.day} de {MESES_ES[hoy.month - 1]} de {hoy.year}"


def pendiente(descripcion: str) -> str:
    return f"[PENDIENTE: {descripcion}]"


def texto_plano_metadata(texto: str) -> str:
    """Version minima para el corchete opcional de \\APMtitleEN (pdftitle):
    con ese corchete vacio, \\ifx\\@apmtitleMeta\\empty deja pdftitle en
    blanco pese a que .cls documenta el fallback a \\@apmtitleEN (verificado
    compilando: vease reporte_tecnico.md) — asi que SIEMPRE se rellena con
    una version plana del titulo. Solo quita backslash/llaves (rompen el
    grupo de \\hypersetup); el resto, tal cual, sin escapar comandos LaTeX,
    igual que taller/ejemplo_editorial.tex."""
    if not texto:
        return ""
    return texto.replace("\\", "").replace("{", "").replace("}", "")


# ═══════════════════════════════════════════════════════════════
# Runs de un parrafo -> LaTeX (preserva negrita/cursiva de Word)
# ═══════════════════════════════════════════════════════════════

def parrafo_a_latex(paragraph: "Paragraph") -> str:
    piezas = []
    for run in paragraph.runs:
        texto = run.text
        if not texto:
            continue
        frag = escapar_latex(texto)
        if run.bold and run.italic:
            frag = r"\textbf{\textit{" + frag + "}}"
        elif run.bold:
            frag = r"\textbf{" + frag + "}"
        elif run.italic:
            frag = r"\textit{" + frag + "}"
        piezas.append(frag)
    if piezas:
        return "".join(piezas)
    # runs vacios (p. ej. solo imagen): recurre al texto plano del parrafo
    return escapar_latex(paragraph.text)


def celda_a_latex(cell) -> str:
    partes = [parrafo_a_latex(p) for p in cell.paragraphs if p.text.strip()]
    return r" \newline ".join(partes) if partes else ""


RE_ROTULO_PIE = re.compile(
    r"^(tabla|cuadro|figura|imagen)\s*\.?\s*\d*\s*[:.\-–]?\s*", re.IGNORECASE
)


def quita_rotulo_pie(texto: str) -> str:
    """Retira el rotulo 'Tabla N.'/'Cuadro N.'/'Figura N.'/'Imagen N.' inicial
    del parrafo-pie: \\caption{} ya antepone ese mismo rotulo y su numero
    (via el contador del entorno y captionsspanish), asi que dejarlo tambien
    en el texto duplicaria "Tabla 1. Tabla 1. ...". El resto del pie, tal
    como lo escribio el autor, no se toca."""
    return RE_ROTULO_PIE.sub("", texto.strip(), count=1)


def referencia_a_latex(texto: str) -> str:
    """Envuelve URLs en \\url{...} (verbatim, sin escapar) para permitir el
    corte de linea del paquete url; el resto se escapa. No reescribe, no
    reordena, no reformatea la cita: mismos caracteres, solo tipografia."""
    piezas = []
    ultimo = 0
    for m in RE_URL.finditer(texto):
        inicio, fin = m.span()
        # no incluir puntuacion de cierre de frase pegada al final de la URL
        url = m.group(0)
        while url and url[-1] in ".,;:)]”’":
            url = url[:-1]
            fin -= 1
        piezas.append(escapar_latex(texto[ultimo:inicio]))
        piezas.append(r"\url{" + url + "}")
        ultimo = fin
    piezas.append(escapar_latex(texto[ultimo:]))
    return "".join(piezas)


# ═══════════════════════════════════════════════════════════════
# Recorrido del documento EN ORDEN (body/parrafos/tablas intercalados)
# ═══════════════════════════════════════════════════════════════

def bloques_en_orden(document):
    """Lista de ('p', Paragraph) / ('tbl', Table) en el orden real del
    documento. Iterar document.paragraphs y document.tables por separado
    pierde el intercalado real: no lo hacemos."""
    bloques = []
    for hijo in document.element.body.iterchildren():
        if hijo.tag == qn("w:p"):
            bloques.append(("p", Paragraph(hijo, document)))
        elif hijo.tag == qn("w:tbl"):
            bloques.append(("tbl", Table(hijo, document)))
    return bloques


def nivel_de_encabezado(paragraph: "Paragraph"):
    """None si no es encabezado; 0 para Title; 1..N para Heading N."""
    try:
        nombre = normaliza(paragraph.style.name)
    except Exception:
        return None
    if nombre == "title":
        return 0
    m = re.match(r"heading\s*(\d+)", nombre)
    if m:
        return int(m.group(1))
    return None


def imagenes_en_parrafo(paragraph: "Paragraph", document):
    """Resuelve los w:drawing/r:embed dentro de cada run, en orden real:
    document.inline_shapes no basta por si solo para preservar la posicion.
    Devuelve lista de (image_part, ancho_pulgadas_o_None)."""
    resultados = []
    for run in paragraph.runs:
        for blip in run._element.findall(f".//{{{A_NS}}}blip"):
            rid = blip.get(f"{{{R_NS}}}embed")
            if not rid:
                continue
            try:
                part = document.part.related_parts[rid]
            except KeyError:
                continue
            ancho_in = None
            extent = run._element.find(f".//{{{WP_NS}}}extent")
            if extent is not None:
                cx = extent.get("cx")
                if cx:
                    ancho_in = int(cx) / 914400.0
            resultados.append((part, ancho_in))
    return resultados


# ═══════════════════════════════════════════════════════════════
# Extraccion del manuscrito
# ═══════════════════════════════════════════════════════════════

class Extraccion:
    def __init__(self):
        self.titulo = None
        self.autores = None
        self.afiliacion = None
        self.email = None
        self.orcid = None
        self.telefono = None
        self.doi = None
        self.resumen_runs = None  # lista de Paragraph
        self.palabras_clave = None
        self.conflicto = None
        self.financiamiento = None
        self.cuerpo_bloques = []  # lista de ('p'|'tbl', obj)
        self.referencias = []  # lista de str (texto plano)
        self.encabezados_vistos = []  # texto normalizado de TODOS los headings
        self.notas = []  # decisiones/heuristicas para reporte_tecnico.md
        self.placeholders = []


def es_heading_resumen(texto_norm):
    return "resumen" in texto_norm or "abstract" in texto_norm


def es_heading_referencias(texto_norm):
    return ("referencias" in texto_norm or "bibliografia" in texto_norm
            or "references" in texto_norm)


def es_heading_conflicto(texto_norm):
    return "conflicto de intereses" in texto_norm or "conflicts of interest" in texto_norm


def es_heading_financiamiento(texto_norm):
    return "financiamiento" in texto_norm or "funding" in texto_norm


def es_palabras_clave(texto_norm):
    return texto_norm.startswith("palabras clave") or texto_norm.startswith("keywords")


MAX_LINEAS_AFILIACION = 3  # mas alla de esto, ya no es un bloque de contacto


def parece_cuerpo(texto):
    """Respaldo para manuscritos SIN ningun encabezado de seccion: una linea
    de afiliacion/contacto es corta y rara vez trae mas de una oracion
    completa. Un parrafo largo, o con dos o mas oraciones, se trata como
    inicio del cuerpo en vez de seguir barriendolo como metadato de
    contacto -- evita que un manuscrito sin heading "Resumen" ni ningun
    otro encabezado pierda su cuerpo dentro de \\APMaffiliation{}."""
    if len(texto) > 280:
        return True
    return len(re.findall(r'[.!?]\s+[A-ZÁÉÍÓÚÑ]', texto)) >= 2


def extraer(document) -> Extraccion:
    ex = Extraccion()
    bloques = bloques_en_orden(document)

    # texto plano completo, para DOI global
    texto_completo = "\n".join(
        p.text for tipo, p in bloques if tipo == "p"
    )
    m = RE_DOI_URL.search(texto_completo) or RE_DOI_BARE.search(texto_completo)
    if m:
        ex.doi = m.group(0).rstrip(").,;")

    # ── Titulo: primer parrafo Title/Heading1, si no el primer no vacio ──
    idx_titulo = None
    for i, (tipo, obj) in enumerate(bloques):
        if tipo != "p":
            continue
        if not obj.text.strip():
            continue
        nivel = nivel_de_encabezado(obj)
        if nivel in (0, 1):
            idx_titulo = i
            break
    if idx_titulo is None:
        for i, (tipo, obj) in enumerate(bloques):
            if tipo == "p" and obj.text.strip():
                idx_titulo = i
                break
    if idx_titulo is not None:
        ex.titulo = bloques[idx_titulo][1].text.strip()
    else:
        ex.titulo = None
        ex.placeholders.append("titulo")

    # ── Autores: primer parrafo no vacio tras el titulo, si no es un
    # encabezado de resumen/referencias/palabras clave ──
    idx_autores = None
    cursor = (idx_titulo + 1) if idx_titulo is not None else 0
    for i in range(cursor, len(bloques)):
        tipo, obj = bloques[i]
        if tipo != "p" or not obj.text.strip():
            continue
        norm = normaliza(obj.text)
        if es_heading_resumen(norm) or es_heading_referencias(norm) or es_palabras_clave(norm):
            break
        idx_autores = i
        ex.autores = obj.text.strip()
        break

    # ── Metadatos de contacto: entre autores y resumen/palabras-clave ──
    #
    # Este barrido NO tiene un limite natural cuando el manuscrito no trae
    # heading "Resumen" (legitimo en Editorial/Carta al Editor, que la norma
    # exime de resumen): sin ese tope, se comia el documento entero como
    # afiliacion y el cuerpo quedaba vacio. Se le ponen dos topes de
    # seguridad, independientes de si aparece un heading de resumen:
    # (1) cualquier OTRO encabezado (Introduccion, Metodo, ...) tambien
    #     cierra el bloque de contacto -- una vez que el autor tituló una
    #     seccion, ya no es metadato de contacto;
    # (2) sin ningun encabezado en absoluto, un parrafo que "parece cuerpo"
    #     (largo, o con varias oraciones) o el cupo maximo de lineas de
    #     afiliacion cierran el bloque igual.
    afiliacion_lineas = []
    idx_resumen_heading = None
    idx_palabras_clave = None
    cursor = (idx_autores + 1) if idx_autores is not None else cursor
    fin_metadatos = len(bloques)
    for i in range(cursor, len(bloques)):
        tipo, obj = bloques[i]
        if tipo != "p":
            continue
        texto = obj.text.strip()
        norm = normaliza(texto)
        es_head = bool(texto) and (nivel_de_encabezado(obj) is not None)
        if es_head and es_heading_resumen(norm):
            idx_resumen_heading = i
            fin_metadatos = i
            break
        if texto and es_palabras_clave(norm):
            idx_palabras_clave = i
            fin_metadatos = i
            break
        if texto and es_heading_referencias(norm):
            fin_metadatos = i
            break
        if es_head:
            # Cualquier otro encabezado (p. ej. "Introduccion" en un
            # Editorial sin Resumen) marca el fin de los metadatos de
            # contacto: el cuerpo empieza aqui.
            fin_metadatos = i
            break
        if not texto:
            continue
        email_m = RE_EMAIL.search(texto)
        orcid_m = RE_ORCID.search(texto)
        es_tel = RE_TEL_CONTEXTO.search(texto) and RE_TEL_NUM.search(texto)
        if email_m and ex.email is None:
            ex.email = email_m.group(0)
            resto = (texto[:email_m.start()] + texto[email_m.end():]).strip(" ,;:•")
            if resto:
                afiliacion_lineas.append(resto)
            continue
        if orcid_m and ex.orcid is None:
            ex.orcid = "https://orcid.org/" + orcid_m.group(0)
            continue
        if es_tel and ex.telefono is None:
            ex.telefono = RE_TEL_NUM.search(texto).group(0).strip()
            continue
        if len(afiliacion_lineas) >= MAX_LINEAS_AFILIACION or parece_cuerpo(texto):
            # Manuscrito sin ningun encabezado: nada distingue "afiliacion"
            # de "cuerpo" salvo la forma del parrafo. Este ya no califica
            # como linea de contacto -- se deja para el cuerpo.
            fin_metadatos = i
            break
        afiliacion_lineas.append(texto)
    if afiliacion_lineas:
        ex.afiliacion = "\n".join(afiliacion_lineas)

    # ── Resumen ──
    if idx_resumen_heading is not None:
        ex.resumen_runs = []
        for i in range(idx_resumen_heading + 1, len(bloques)):
            tipo, obj = bloques[i]
            if tipo != "p":
                break
            texto = obj.text.strip()
            norm = normaliza(texto)
            if texto and es_palabras_clave(norm):
                idx_palabras_clave = i
                break
            if texto and nivel_de_encabezado(obj) is not None:
                break
            if texto and es_heading_referencias(norm):
                break
            if texto:
                ex.resumen_runs.append(obj)
            idx_resumen_heading = i
        inicio_cuerpo = i if idx_palabras_clave is None else idx_palabras_clave
    else:
        inicio_cuerpo = fin_metadatos

    # ── Palabras clave ──
    if idx_palabras_clave is not None:
        texto = bloques[idx_palabras_clave][1].text.strip()
        etiqueta_m = re.match(r"(?i)^(palabras\s+clave|keywords)\s*[:\-]?\s*(.*)$", texto)
        contenido = etiqueta_m.group(2) if etiqueta_m else texto
        crudas = re.split(r"[;,]", contenido)
        ex.palabras_clave = [c.strip() for c in crudas if c.strip()]
        inicio_cuerpo = idx_palabras_clave + 1

    # ── Referencias / Conflicto / Financiamiento: localizar encabezados ──
    idx_referencias = None
    idx_conflicto = None
    idx_financiamiento = None
    for i, (tipo, obj) in enumerate(bloques):
        if tipo != "p" or not obj.text.strip():
            continue
        norm = normaliza(obj.text)
        es_head = nivel_de_encabezado(obj) is not None
        if es_head:
            ex.encabezados_vistos.append(norm)
        if es_head and idx_referencias is None and es_heading_referencias(norm):
            idx_referencias = i
        if es_head and idx_conflicto is None and es_heading_conflicto(norm):
            idx_conflicto = i
        if es_head and idx_financiamiento is None and es_heading_financiamiento(norm):
            idx_financiamiento = i

    fin_cuerpo = idx_referencias if idx_referencias is not None else len(bloques)
    if idx_conflicto is not None:
        fin_cuerpo = min(fin_cuerpo, idx_conflicto)
    if idx_financiamiento is not None:
        fin_cuerpo = min(fin_cuerpo, idx_financiamiento)

    # conflicto/financiamiento como heading dedicado (parrafo siguiente)
    if idx_conflicto is not None:
        siguientes = []
        for i in range(idx_conflicto + 1, len(bloques)):
            tipo, obj = bloques[i]
            if tipo != "p":
                break
            if obj.text.strip() and nivel_de_encabezado(obj) is not None:
                break
            if obj.text.strip():
                siguientes.append(obj.text.strip())
            else:
                if siguientes:
                    break
        if siguientes:
            ex.conflicto = " ".join(siguientes)
    if idx_financiamiento is not None:
        siguientes = []
        for i in range(idx_financiamiento + 1, len(bloques)):
            tipo, obj = bloques[i]
            if tipo != "p":
                break
            if obj.text.strip() and nivel_de_encabezado(obj) is not None:
                break
            if obj.text.strip():
                siguientes.append(obj.text.strip())
            else:
                if siguientes:
                    break
        if siguientes:
            ex.financiamiento = " ".join(siguientes)

    # ── Cuerpo: bloques entre inicio_cuerpo y fin_cuerpo, excluyendo los
    # parrafos-etiqueta de conflicto/financiamiento en linea ──
    excluidos = set()
    for i in range(inicio_cuerpo, fin_cuerpo):
        tipo, obj = bloques[i]
        if tipo != "p":
            continue
        norm = normaliza(obj.text)
        if ex.conflicto is None and norm.startswith("conflicto de intereses"):
            partes = re.split(r"[:\-]", obj.text, maxsplit=1)
            if len(partes) == 2 and partes[1].strip():
                ex.conflicto = partes[1].strip()
                excluidos.add(i)
        if ex.financiamiento is None and (norm.startswith("financiamiento") or norm.startswith("funding")):
            partes = re.split(r"[:\-]", obj.text, maxsplit=1)
            if len(partes) == 2 and partes[1].strip():
                ex.financiamiento = partes[1].strip()
                excluidos.add(i)

    ex.cuerpo_bloques = [
        bloques[i] for i in range(inicio_cuerpo, fin_cuerpo) if i not in excluidos
    ]

    # ── Referencias: cada parrafo no vacio tras el heading, literal ──
    if idx_referencias is not None:
        limite = len(bloques)
        for candidato in (idx_conflicto, idx_financiamiento):
            if candidato is not None and candidato > idx_referencias:
                limite = min(limite, candidato)
        for i in range(idx_referencias + 1, limite):
            tipo, obj = bloques[i]
            if tipo != "p":
                continue
            texto = obj.text.strip()
            if texto:
                ex.referencias.append(texto)

    return ex


# ═══════════════════════════════════════════════════════════════
# Heuristica de \APMtype
# ═══════════════════════════════════════════════════════════════

def decide_tipo(ex: Extraccion, titulo_norm: str, n_palabras_cuerpo: int, tipo_forzado):
    if tipo_forzado:
        return tipo_forzado, "regla 1: --tipo dado explicitamente"
    headings = ex.encabezados_vistos
    tiene_metodo = any(("metodo" in h or "metodologia" in h) for h in headings)
    tiene_resultados = any("resultados" in h for h in headings)
    if tiene_metodo and tiene_resultados:
        return "Artículo original", "regla 2: encabezados de metodo Y resultados"
    universo = headings + [titulo_norm]
    if any(("caso clinico" in h or "presentacion de caso" in h) for h in universo):
        return "Caso clínico", "regla 3: titulo/encabezados mencionan caso clinico"
    if any(("revision sistematica" in h or "meta-analisis" in h or "meta analisis" in h or "revision" in h)
           for h in universo):
        return "Artículo de revisión", "regla 4: titulo/encabezados mencionan revision/meta-analisis"
    if n_palabras_cuerpo < 600:
        return "Carta al Editor", f"regla 5: cuerpo corto ({n_palabras_cuerpo} palabras) y ninguna regla previa aplico"
    return "Artículo original", "regla 6: ninguna regla previa aplico (caso mas comun)"


# ═══════════════════════════════════════════════════════════════
# Volumen / numero / periodo / ART#
# ═══════════════════════════════════════════════════════════════

PERIODOS = {1: "ENERO--ABRIL", 2: "MAYO--AGOSTO", 3: "SEPTIEMBRE--DICIEMBRE"}


def numero_por_omision():
    hoy = datetime.date.today()
    vol = hoy.year - 2020
    if hoy.month <= 4:
        no = 1
    elif hoy.month <= 8:
        no = 2
    else:
        no = 3
    return f"VOL{vol}_NO{no}"


def parsea_numero(numero_str):
    m = re.match(r"(?i)^VOL(\d+)_NO(\d+)$", numero_str.strip())
    if not m:
        raise ValueError(f"--numero mal formado: {numero_str!r} (se espera VOLx_NOy)")
    vol = int(m.group(1))
    no = int(m.group(2))
    if no not in PERIODOS:
        raise ValueError(f"--numero con periodo invalido: NO{no} (debe ser 1, 2 o 3)")
    anio = vol + 2020
    return vol, no, anio, PERIODOS[no]


def siguiente_art(raiz_numero_dir, numero_norm):
    if not os.path.isdir(raiz_numero_dir):
        return 1
    patron = re.compile(rf"^{re.escape(numero_norm)}_ART(\d+)_")
    maximo = 0
    for nombre in os.listdir(raiz_numero_dir):
        m = patron.match(nombre)
        if m:
            maximo = max(maximo, int(m.group(1)))
    return maximo + 1


# ═══════════════════════════════════════════════════════════════
# Render de cuerpo (encabezados, tablas, figuras, parrafos)
# ═══════════════════════════════════════════════════════════════

def mapa_niveles_encabezado(cuerpo_bloques):
    """El nivel de heading MAS ALTO usado de forma consistente -> \\section*;
    el siguiente nivel -> \\subsection*. Niveles mas profundos se funden con
    el segundo (declarado en el reporte si ocurre)."""
    niveles = sorted({
        nivel_de_encabezado(obj) for tipo, obj in cuerpo_bloques
        if tipo == "p" and obj.text.strip() and nivel_de_encabezado(obj) not in (None, 0)
    })
    mapa = {}
    if len(niveles) >= 1:
        mapa[niveles[0]] = r"\section*"
    for n in niveles[1:]:
        mapa[n] = r"\subsection*"
    return mapa, niveles


TABCOLSEP_CM = 6 / 72.27 * 2.54  # \tabcolsep por omision (6pt), a cada lado de cada columna


def renderiza_tabla(table: "Table", caption_tex, indice, notas):
    filas = table.rows
    ncols = max((len(r.cells) for r in filas), default=0)
    ancho_pagina = ncols >= 4
    entorno = "table*" if ancho_pagina else "table"
    # textwidth ~18.0cm (margenes 1.8cm); columna sencilla ~8.75cm
    # (columnsep 5mm). Se resta el \tabcolsep de cada lado de cada columna
    # y un margen de seguridad para no producir Overfull.
    disponible_cm = 17.4 if ancho_pagina else 8.3
    disponible_cm -= 2 * ncols * TABCOLSEP_CM
    colw = max(1.2, disponible_cm / max(ncols, 1))
    colspec = " ".join([f"p{{{colw:.2f}cm}}"] * ncols)

    lineas = [rf"\begin{{{entorno}}}[t]", r"\centering",
              rf"\caption{{{caption_tex}}}", rf"\label{{tab:{indice}}}",
              rf"\begin{{tabular}}{{{colspec}}}", r"\toprule"]
    for i, fila in enumerate(filas):
        celdas = [celda_a_latex(c) for c in fila.cells[:ncols]]
        while len(celdas) < ncols:
            celdas.append("")
        if i == 0:
            celdas = [rf"\textbf{{{c}}}" if c else c for c in celdas]
        lineas.append(" & ".join(celdas) + r" \\")
        if i == 0:
            lineas.append(r"\midrule")
    lineas += [r"\bottomrule", r"\end{tabular}", rf"\end{{{entorno}}}"]
    return "\n".join(lineas)


def renderiza_figura(image_part, ancho_in, caption_tex, indice, medios_dir, nombre_tex_relativo, notas):
    ext = os.path.splitext(image_part.partname)[1] or ".png"
    nombre_archivo = f"fig{indice}{ext}"
    with open(os.path.join(medios_dir, nombre_archivo), "wb") as fh:
        fh.write(image_part.blob)
    ancho_pagina = bool(ancho_in and ancho_in > 3.4)
    entorno = "figure*" if ancho_pagina else "figure"
    ancho_incl = "0.9\\linewidth" if not ancho_pagina else "0.55\\textwidth"
    ruta = f"{nombre_tex_relativo}/{nombre_archivo}"
    return "\n".join([
        rf"\begin{{{entorno}}}[t]", r"\centering",
        rf"\includegraphics[width={ancho_incl}]{{{ruta}}}",
        rf"\caption{{{caption_tex}}}", rf"\label{{fig:{indice}}}",
        rf"\end{{{entorno}}}",
    ])


def renderiza_cuerpo(ex: Extraccion, document, medios_dir, notas):
    notas.append(
        "pies de tabla/figura: cuando el parrafo precedente empieza con "
        "'Tabla N.'/'Cuadro N.'/'Figura N.'/'Imagen N.', ese rotulo y numero "
        "se retiran del texto del pie antes de pasarlo a \\caption{}, porque "
        "\\caption ya antepone 'Tabla N:'/'Figura N:' automaticamente "
        "(babel captionsspanish) y dejarlo tambien en el texto duplicaria "
        "el rotulo; el resto del pie, tal como lo escribio el autor, no se "
        "toca."
    )
    mapa, niveles = mapa_niveles_encabezado(ex.cuerpo_bloques)
    if len(niveles) > 2:
        notas.append(
            f"el manuscrito uso {len(niveles)} niveles de encabezado en el cuerpo "
            f"({niveles}); el primero se rotulo \\section*, TODOS los siguientes "
            f"se fundieron en \\subsection* (el .cls no se toco para agregar un "
            f"tercer nivel automatico)."
        )

    bloques = ex.cuerpo_bloques
    salida = []
    idx_tabla = 0
    idx_figura = 0
    i = 0
    while i < len(bloques):
        tipo, obj = bloques[i]
        if tipo == "tbl":
            idx_tabla += 1
            caption = None
            # buscar el parrafo INMEDIATO anterior ya emitido como candidato
            if salida and salida[-1][0] == "caption_candidata_tabla":
                caption = salida.pop()[1]
            if caption is None:
                caption = pendiente(f"pie de tabla {idx_tabla}")
                notas.append(f"tabla {idx_tabla}: sin parrafo 'Tabla/Cuadro' precedente; pie generico pendiente para el editor.")
            salida.append(("tex", renderiza_tabla(obj, caption, idx_tabla, notas)))
            i += 1
            continue
        if tipo == "p":
            texto = obj.text.strip()
            imagenes = imagenes_en_parrafo(obj, document) if texto or obj.runs else []
            if imagenes:
                for parte, ancho_in in imagenes:
                    idx_figura += 1
                    caption = None
                    if salida and salida[-1][0] == "caption_candidata_figura":
                        caption = salida.pop()[1]
                    if caption is None:
                        caption = pendiente(f"pie de figura {idx_figura}")
                        notas.append(f"figura {idx_figura}: sin parrafo 'Figura/Imagen' precedente; pie generico pendiente para el editor.")
                    salida.append(("tex", renderiza_figura(
                        parte, ancho_in, caption, idx_figura, medios_dir, "medios", notas)))
                i += 1
                continue
            if not texto:
                i += 1
                continue
            norm = normaliza(texto)
            nivel = nivel_de_encabezado(obj)
            # candidata a pie de tabla/figura: se decide al llegar al flotante
            if norm.startswith("tabla") or norm.startswith("cuadro"):
                sig_tipo, sig_obj = bloques[i + 1] if i + 1 < len(bloques) else (None, None)
                if sig_tipo == "tbl":
                    salida.append(("caption_candidata_tabla", escapar_latex(quita_rotulo_pie(texto))))
                    i += 1
                    continue
            if norm.startswith("figura") or norm.startswith("imagen"):
                if i + 1 < len(bloques):
                    sig_tipo, sig_obj = bloques[i + 1]
                    if sig_tipo == "p" and imagenes_en_parrafo(sig_obj, document):
                        salida.append(("caption_candidata_figura", escapar_latex(quita_rotulo_pie(texto))))
                        i += 1
                        continue
            if nivel is not None and nivel in mapa:
                salida.append(("tex", f"{mapa[nivel]}{{{escapar_latex(texto)}}}\n"))
            else:
                salida.append(("tex", parrafo_a_latex(obj)))
            i += 1
            continue
        i += 1
    piezas = [contenido for tipo, contenido in salida if tipo == "tex"]
    return "\n\n".join(piezas)


# ═══════════════════════════════════════════════════════════════
# Generacion del .tex
# ═══════════════════════════════════════════════════════════════

TEX_PLANTILLA = r"""% ═══════════════════════════════════════════════════════════════
% Generado por taller/recibir_articulo.py — NO editar a mano el
% preambulo sin volver a correr el receptor; el cuerpo si es editable.
% Manuscrito origen: {origen}
% ═══════════════════════════════════════════════════════════════
\documentclass{{apm-editorial}}

% ── METADATOS ────────────────────────────────────────────────
\APMtype{{{tipo}}}
\APMtitle{{{titulo}}}
\APMtitleEN[{titulo_meta}]{{{titulo}}}
\APMauthor{{{autores}}}
\APMaffiliation{{{afiliacion}}}
\APMorcid{{{orcid}}}
\APMemail{{{email}}}
\APMphone{{{telefono}}}
\APMaddress{{{direccion}}}
\APMdates{{{fecha_recibido}}}{{{fecha_aceptado}}}{{{fecha_publicado}}}
\APMdoi{{{doi}}}
\APMauthorshort{{{autor_corto}}}
\APMauthorpdf{{{autores_pdf}}}
\APMkeywordspdf{{{keywords_pdf}}}
\APMtitleshort{{{titulo_corto}}}
\APMvolume{{{vol}}}{{{num}}}{{{anio}}}
\APMperiod{{{periodo}}}
\APMlogo{{logo_hires.png}}

% ── RESUMEN ──────────────────────────────────────────────────
\APMabstract{{{resumen}}}
\APMkeywords{{{keywords}}}

% ── COLOFÓN ──────────────────────────────────────────────────
\APMconflict{{{conflicto}}}
\APMfunding{{{financiamiento}}}
\APMlicense{{\textcopyright\ {anio} Los Autores. Publicado bajo Creative
  Commons Atribución-No Comercial 4.0 Internacional (CC~BY-NC~4.0).}}

% ═════════════════════════════════════════════════════════════
\begin{{document}}
\makeAPMeditorial
\pdfinfo{{
  /Author ({autor_octal})
  /Keywords ({keywords_octal})
}}

% ── CUERPO ───────────────────────────────────────────────────

{cuerpo}

% ── REFERENCIAS ──────────────────────────────────────────────
\APMrefsstart

{referencias}

\APMrefsend

% ── COLOFÓN ──────────────────────────────────────────────────
{{\hfuzz=15pt\APMcolophon}}

\end{{document}}
"""


def construir_tex(ex: Extraccion, document, tipo, vol, num, anio, periodo,
                   medios_dir, notas, origen_nombre):
    titulo = ex.titulo or pendiente("titulo")
    if ex.titulo is None:
        notas.append("titulo: no se encontro ningun parrafo con texto en el manuscrito; se inserto placeholder.")

    autores = ex.autores or pendiente("autor(es)")
    if ex.autores is None:
        notas.append("autores: no se encontro un parrafo candidato tras el titulo; se inserto placeholder.")

    afiliacion = ex.afiliacion or pendiente("filiación institucional")
    if ex.afiliacion is None:
        notas.append("afiliacion: no se encontro ningun parrafo de contacto entre autores y resumen; placeholder.")

    orcid = ex.orcid or pendiente("ORCID")
    if ex.orcid is None:
        notas.append("ORCID: no se encontro un identificador con el patron NNNN-NNNN-NNNN-NNN[X]; placeholder.")

    email = ex.email or pendiente("correo de contacto")
    if ex.email is None:
        notas.append("correo: no se encontro una direccion de correo en los parrafos de contacto; placeholder.")

    telefono = ex.telefono or pendiente("teléfono de contacto")
    if ex.telefono is None:
        notas.append("telefono: no se encontro un numero cerca de 'tel'/'teléfono'/'phone'; placeholder.")

    direccion = pendiente("domicilio de correspondencia")
    notas.append("domicilio: el manuscrito .docx no distingue domicilio postal de filiacion institucional; se dejo pendiente para el editor (no se infiere del texto de afiliacion para no inventar un dato distinto).")

    doi = ex.doi or "[DOI pendiente de asignación]"
    if ex.doi is None:
        notas.append("DOI: no se encontro un DOI ni una URL doi.org en el manuscrito; se uso el placeholder de asignacion editorial.")

    if ex.resumen_runs:
        resumen = " ".join(parrafo_a_latex(p) for p in ex.resumen_runs)
    else:
        resumen = pendiente("resumen no incluido en el manuscrito")
        notas.append("resumen: no se encontro un encabezado 'Resumen'/'Abstract'; se inserto placeholder (la caja de resumen se imprime igual, marcada como pendiente).")

    if ex.palabras_clave:
        keywords = ", ".join(escapar_latex(k) for k in ex.palabras_clave)
        keywords_plano = ", ".join(ex.palabras_clave)
    else:
        keywords = pendiente("palabras clave")
        keywords_plano = "PENDIENTE: palabras clave"
        notas.append("palabras clave: no se encontro un parrafo 'Palabras clave:'/'Keywords:'; se inserto placeholder.")

    conflicto = ex.conflicto or "[PENDIENTE: declarar conflicto de intereses]"
    if ex.conflicto is None:
        notas.append("conflicto de intereses: no declarado en el manuscrito; placeholder exacto segun el contrato.")
    financiamiento = ex.financiamiento or "[PENDIENTE: declarar financiamiento]"
    if ex.financiamiento is None:
        notas.append("financiamiento: no declarado en el manuscrito; placeholder exacto segun el contrato.")

    autores_split = re.split(r",| y | and ", ex.autores) if ex.autores else []
    primer_autor = autores_split[0].strip() if autores_split and autores_split[0].strip() else (ex.autores or "Autor")
    autor_corto = primer_autor + (" et al." if len(autores_split) > 1 else "")
    if len(autores_split) > 1:
        notas.append("autorshort: se tomo el primer autor listado + 'et al.' (no se intento separar nombre/apellido: el docx no distingue campos).")

    titulo_corto = titulo if len(titulo) <= 70 else titulo[:67].rsplit(" ", 1)[0] + "..."

    cuerpo = renderiza_cuerpo(ex, document, medios_dir, notas)
    referencias = "\n\n".join(rf"\APMref{{{referencia_a_latex(r)}}}" for r in ex.referencias)
    if not referencias:
        referencias = rf"\APMref{{{pendiente('lista de referencias no incluida en el manuscrito')}}}"
        notas.append("referencias: no se encontro un encabezado 'Referencias/Bibliografía/References'; se inserto un placeholder unico.")

    tex = TEX_PLANTILLA.format(
        origen=origen_nombre,
        tipo=escapar_latex(tipo),
        titulo=escapar_latex(titulo),
        titulo_meta=texto_plano_metadata(titulo),
        autores=escapar_latex(autores),
        afiliacion=escapar_latex(afiliacion),
        orcid=escapar_latex(orcid),
        email=escapar_latex(email),
        telefono=escapar_latex(telefono),
        direccion=escapar_latex(direccion),
        fecha_recibido=fecha_hoy_es(),
        fecha_aceptado="[PENDIENTE]",
        fecha_publicado="[PENDIENTE]",
        doi=escapar_latex(doi),
        autor_corto=escapar_latex(autor_corto),
        autores_pdf=autores,
        keywords_pdf=keywords_plano,
        titulo_corto=escapar_latex(titulo_corto),
        vol=vol, num=num, anio=anio, periodo=periodo,
        resumen=resumen,
        keywords=keywords,
        conflicto=escapar_latex(conflicto),
        financiamiento=escapar_latex(financiamiento),
        autor_octal=a_octal_pdfinfo(autores),
        keywords_octal=a_octal_pdfinfo(keywords_plano),
        cuerpo=cuerpo,
        referencias=referencias,
    )
    return tex


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Recibe un manuscrito .docx y produce el articulo camera-ready de la Revista APM.")
    ap.add_argument("manuscrito", help="ruta al MANUSCRITO.docx")
    ap.add_argument("--tipo", default=None, help='fuerza \\APMtype, p.ej. "Artículo original"')
    ap.add_argument("--numero", default=None, help="carpeta de destino, p.ej. VOL6_NO3 (por omision, calculada de hoy)")
    ap.add_argument("--raiz", default=None, help='raiz de numeros/ (por omision "numeros/" desde la raiz del repo)')
    ap.add_argument("--art", type=int, default=None, help="ART# a usar (por omision, el siguiente libre)")
    args = ap.parse_args(argv)

    if not os.path.isfile(args.manuscrito):
        sys.exit(f"no existe el manuscrito: {args.manuscrito}")

    numero_str = args.numero or numero_por_omision()
    try:
        vol, num, anio, periodo = parsea_numero(numero_str)
    except ValueError as exc:
        sys.exit(str(exc))
    numero_norm = f"VOL{vol}_NO{num}"

    raiz = args.raiz or os.path.join(RAIZ_REPO, "numeros")
    raiz = os.path.abspath(raiz)
    os.makedirs(raiz, exist_ok=True)

    raiz_numero_dir = os.path.join(raiz, numero_norm)
    art_num = args.art or siguiente_art(raiz_numero_dir, numero_norm)

    document = docx.Document(args.manuscrito)
    ex = extraer(document)

    if not ex.cuerpo_bloques:
        # Nunca produce un PDF "camera-ready" con el cuerpo vacio: eso se ve
        # terminado y no lo esta. Pasa solo en un manuscrito sin NINGUN
        # encabezado de seccion y con un cuerpo demasiado corto para que la
        # heuristica de metadatos de contacto (ver extraer()) lo distinga de
        # una linea de afiliacion -- fuera de ese extremo, no ocurre.
        sys.exit(
            "el receptor no encontro texto de cuerpo: el manuscrito no trae "
            "ningun encabezado de seccion (ni siquiera \"Introduccion\") y su "
            "parrafo de cuerpo es demasiado corto para distinguirlo de una "
            "linea de afiliacion de autor. Anade un encabezado de seccion "
            "antes del texto del cuerpo, o alarga ese parrafo, y vuelve a "
            "intentarlo -- no se genera un PDF con el cuerpo en blanco."
        )

    titulo_para_tipo = ex.titulo or ""
    # palabras del cuerpo (solo texto de parrafos normales, sin tablas/figuras)
    n_palabras = sum(
        len(obj.text.split()) for tipo, obj in ex.cuerpo_bloques if tipo == "p"
    )
    tipo, regla = decide_tipo(ex, normaliza(titulo_para_tipo), n_palabras, args.tipo)
    ex.notas.append(f"\\APMtype elegido: \"{tipo}\" ({regla}).")

    primera = primera_palabra_significativa(ex.titulo or "ARTICULO")
    carpeta_articulo = f"{numero_norm}_ART{art_num}_{primera}"
    destino = os.path.join(raiz_numero_dir, carpeta_articulo)
    os.makedirs(destino, exist_ok=True)
    medios_dir = os.path.join(destino, "medios")
    os.makedirs(medios_dir, exist_ok=True)

    tex_texto = construir_tex(
        ex, document, tipo, vol, num, anio, periodo, medios_dir, ex.notas,
        os.path.basename(args.manuscrito),
    )
    if not os.listdir(medios_dir):
        os.rmdir(medios_dir)

    ruta_tex = os.path.join(destino, f"{primera}.tex")
    with open(ruta_tex, "w", encoding="utf-8") as fh:
        fh.write(tex_texto)

    nombre_pdf_final = f"{primera}_APM_{numero_norm}_{anio}.pdf"
    ruta_pdf_final = os.path.join(destino, nombre_pdf_final)

    componer = os.path.join(AQUI, "componer.sh")
    resultado = subprocess.run(
        ["bash", componer, ruta_tex], capture_output=True, text=True,
    )
    compiló = resultado.returncode == 0
    ruta_pdf_generado = os.path.join(destino, f"{primera}.pdf")
    if compiló and os.path.isfile(ruta_pdf_generado):
        shutil.move(ruta_pdf_generado, ruta_pdf_final)

    reporte = construir_reporte(
        ex, tipo, regla, numero_norm, art_num, vol, num, anio, periodo,
        n_palabras, ruta_tex, ruta_pdf_final if compiló else None,
        resultado, args,
    )
    with open(os.path.join(destino, "reporte_tecnico.md"), "w", encoding="utf-8") as fh:
        fh.write(reporte)

    print(resultado.stdout)
    if resultado.stderr:
        print(resultado.stderr, file=sys.stderr)

    if not compiló:
        print(f"COMPILACION FALLIDA: vease {destino}", file=sys.stderr)
        print(f".tex y reporte_tecnico.md escritos en {destino} para depurar.", file=sys.stderr)
        return 1

    print(f"escrito: {ruta_tex}")
    print(f"escrito: {ruta_pdf_final}")
    print(f"escrito: {os.path.join(destino, 'reporte_tecnico.md')}")
    return 0


def construir_reporte(ex, tipo, regla, numero_norm, art_num, vol, num, anio,
                       periodo, n_palabras, ruta_tex, ruta_pdf, resultado_compilacion, args):
    n_tablas = sum(1 for t, o in ex.cuerpo_bloques if t == "tbl")
    lineas = [
        f"# Reporte técnico — {ex.titulo or '[PENDIENTE: título]'}",
        "",
        "Generado automáticamente por `taller/recibir_articulo.py`. No se",
        "modificó el texto de los autores; los datos ausentes se marcaron",
        "`[PENDIENTE: ...]` y se listan abajo para que el editor los complete.",
        "",
        "## Ubicación",
        f"- Número: {numero_norm} (Vol. {vol}, No. {num}, {anio}, {periodo})",
        f"- ART#: {art_num}",
        f"- .tex: `{ruta_tex}`",
        f"- .pdf: `{ruta_pdf if ruta_pdf else '(no generado: compilación fallida, ver abajo)'}`",
        "",
        "## Tipo de artículo",
        f"- `\\APMtype{{{tipo}}}` — {regla}",
        "",
        "## Estadísticas medidas sobre la extracción",
        f"- Palabras del cuerpo (sin resumen/tablas/figuras/referencias): {n_palabras}",
        f"- Tablas: {n_tablas}",
        f"- Referencias: {len(ex.referencias)}",
        "",
        "## Decisiones y heurísticas declaradas",
    ]
    if ex.notas:
        lineas += [f"- {n}" for n in ex.notas]
    else:
        lineas.append("- (ninguna decisión no trivial)")
    lineas += ["", "## Verificación de compilación (taller/componer.sh)"]
    lineas.append(f"- Código de salida: {resultado_compilacion.returncode}")
    lineas.append("```")
    lineas.append((resultado_compilacion.stdout or "").strip())
    if resultado_compilacion.stderr:
        lineas.append((resultado_compilacion.stderr or "").strip())
    lineas.append("```")
    return "\n".join(lineas) + "\n"


if __name__ == "__main__":
    sys.exit(main())
